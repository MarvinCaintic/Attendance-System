## imports
import mysql
import mysql.connector
import os, sys, sqlite3
from sqlite3 import Error
import FRfDTR
import numpy as np
from PIL import Image
from tkinter import *
from tkinter import filedialog

from PySide2.QtWidgets import QTableWidgetItem
from PySide2.QtGui import *
from PySide2.QtCore import QRunnable, Signal, QObject 

class AppFunctions():
    
    """Docstring for AppFunctions"""
    def __init__(self, arg):
        super(AppFunctions, self).__init__()
        self.arg = arg
        
    ##Create database connection
    def create_connection():
        """ Create a database connection to a SQLite database"""
        mydb = mysql.connector.connect(
            host="localhost",
            user="root",
            password="",
            database="frfdtr_db"
        )
        return mydb
    
    def databaseCreation():
        #create db connection
        """ Create a connection to mysql server"""
        conn = mysql.connector.connect(
            host="localhost",
            user="root",
            password=""
        )
        mycursor = conn.cursor()
        
        ## Create Dabase if not existing
        mycursor.execute("CREATE DATABASE IF NOT EXISTS frfdtr_db")
        
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor()
        
        # Create the employee table first
        mycursor.execute("""CREATE TABLE IF NOT EXISTS employee (
                            id VARCHAR(10),
                            firstName VARCHAR(15),
                            middleName VARCHAR(15),
                            lastName VARCHAR(15),
                            prefix VARCHAR(3),
                            bdate DATE,
                            gender ENUM('male', 'female'),
                            email VARCHAR(30),
                            mobileNumber VARCHAR(11),
                            stats ENUM('active', 'inactive'),
                            PRIMARY KEY(id)
                         )
                         """)
        
        #Next create the table that requires employee_id as foreighn key
        #Create attendance table
        mycursor.execute("""CREATE TABLE IF NOT EXISTS attendance(
                            id INT(11),
                            employee_id VARCHAR(10),
                            date DATE,
                            AM_timeIn TIME,
                            AM_timeOut TIME,
                            PM_timeIn TIME,
                            PM_timeOut TIME,
                            OT_timeIn TIME,
                            OT_timeOut TIME,
                            PRIMARY KEY(id),
                            FOREIGN KEY (employee_id) REFERENCES employee(id)
                        )
                         """)
    
    
    
    def clearFormInput(self):
        #Clear form input
        #Name of the employee
        firstName = self.ui.firstName.setText("")
        middleName = self.ui.middleName.setText("")
        lastName = self.ui.lastName.setText("")
        prefix = self.ui.prefix.setText("")
        
            
        #Contacts and folder address
        email = self.ui.email.setText("")
        mobileNumber = self.ui.mobileNumber.setText("")
        
        
    
    #Get all employee
    def getAllEmployee():
        #create db connection
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor()
        mycursor.execute("SELECT * FROM employee WHERE stats = 'active'")
        myresult = mycursor.fetchall()
        return myresult
    
    def getEmployeeById(id):
        #create db connection
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor()
        mycursor.execute("SELECT * FROM employee WHERE id = %s", id)
        myresult = mycursor.fetchall()
        return myresult[0]
    
    
    def addEmployee(self):
            ##Get form values
        #Name of the employee
        firstName = self.ui.firstName.text()
        middleName = self.ui.middleName.text()
        lastName = self.ui.lastName.text()
        prefix = self.ui.prefix.text()
            
        #Contacts and folder address
        email = self.ui.email.text()
        mobileNumber = self.ui.mobileNumber.text()
        
        # #Other information
        bDate = self.ui.birthDate.date().toString("yyyy-MM-dd")
        
        #Gender
        gender = ""
        if self.ui.female.isChecked():
            gender = "female"
        elif self.ui.male.isChecked():
            gender = "male"
        
        sql = ""
        val = (1, )
        idNumber = self.ui.idNumber.text()
        
        #Create sql statement
        if self.action == "HIRE":
            sql = """INSERT INTO employee (id, firstName, middleName, lastName, prefix, bdate
            , gender, email, mobileNumber,stats) VALUES (%s, %s, %s, %s, %s, %s, %s, %s, %s, 'active');"""
            val = (idNumber, firstName, middleName, lastName, prefix, bDate, gender, email, mobileNumber)
        else:
            row = self.ui.tableWidget.currentRow()
            ID = (self.ui.tableWidget.item(row, 0).text(), )
            sql = """UPDATE employee SET id = %s, firstName = %s, middleName = %s, lastName = %s, prefix = %s
            , bdate = %s, gender = %s, email = %s, mobileNumber = %s WHERE employee.id = %s;"""
            val = (idNumber, firstName, middleName, lastName, prefix, bDate, gender, email, mobileNumber, ID)
        #execute sql statement
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor()
        mycursor.execute(sql, val)
        conn.commit()
            
        #Clear form input
        AppFunctions.clearFormInput(self)
        AppFunctions.displayEmployee(self, AppFunctions.getAllEmployee())
    
    
    def setEmployee(self, action):
        if action == "edit":
            row = self.ui.tableWidget.currentRow()
            id = (self.ui.tableWidget.item(row, 0).text(), )
            employee =  AppFunctions.getEmployeeById(id)
            
            #Name of the employee
            firstName = self.ui.firstName.setText(employee[1])
            middleName = self.ui.middleName.setText(employee[2])
            lastName = self.ui.lastName.setText(employee[3])
            prefix = self.ui.prefix.setText(employee[4])
            
                
            #Contacts and folder address
            email = self.ui.email.setText(employee[7])
            mobileNumber = self.ui.mobileNumber.setText(employee[8])
            
            
            
            
        elif action == "hire":
            #Clear form input
            AppFunctions.clearFormInput(self)
            
    #Display employee
    def displayEmployee(self, rows):
        #Create new row
        self.ui.tableWidget.setRowCount(0)
        for row in rows:
            #Get the number of rows
            rowPosition = self.ui.tableWidget.rowCount()
            if rowPosition  > len(row):
                continue
            
            itemCount = 0
            #Crete new table row
            self.ui.tableWidget.setRowCount(rowPosition + 1)
            qtablewidgetitem = QTableWidgetItem()
            self.ui.tableWidget.setVerticalHeaderItem(rowPosition, qtablewidgetitem)
            myList = [row[0], row[1] + " " + row[2] + " " + row[3] + " " + row[4], row[7], row[8]]
            #add items to row
            for item in myList:
                qtablewidgetitem = QTableWidgetItem()
                self.ui.tableWidget.setItem(rowPosition, itemCount, qtablewidgetitem)
                qtablewidgetitem.setText(str(item))
                itemCount = itemCount+1
            rowPosition = rowPosition+1
            
    def getAttendance(self, dDate):
        #create db connection
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor() 
        d = dDate.toString("yyyy-MM-dd")
        mycursor.execute("SELECT * FROM attendance WHERE date = '" + d + "'")
        myresult = mycursor.fetchall()
        return myresult       
    
    def getAttendanceWithId(id):
        from datetime import date
        #create db connection
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor() 
        today = date.today()
        mycursor.execute("SELECT * FROM attendance WHERE date = %s and employee_id = %s", (today, id))
        myresult = mycursor.fetchall()
        return myresult
    
    def getAttendance2(seld, dDate, id):
        from datetime import date
        #create db connection
        conn = AppFunctions.create_connection()
        mycursor = conn.cursor() 
        d = dDate.toString("yyyy-MM") + "-01"
        mycursor.execute("SELECT * FROM attendance WHERE date BETWEEN %s and last_day(%s) and employee_id = %s order by date", (d, d, id))
        myresult = mycursor.fetchall()
        return myresult
    
    #Display employee
    def displayAttendance(self, codeX):
        if codeX == 1:
            dDate = self.ui.aDate.date()
            rows = AppFunctions.getAttendance(self, dDate)
            #Create new row
            self.ui.attendanceTable.setRowCount(0)
            for row in rows:
                #Get the number of rows
                rowPosition = self.ui.attendanceTable.rowCount()
                if rowPosition  > row[0]:
                    continue
                itemCount = 0
                #Crete new table row
                self.ui.attendanceTable.setRowCount(rowPosition + 1)
                qtablewidgetitem = QTableWidgetItem()
                self.ui.attendanceTable.setVerticalHeaderItem(rowPosition, qtablewidgetitem)
                
                id = (row[1],)
                employee =  AppFunctions.getEmployeeById(id)
                
                #Name of the employee
                empName = str(employee[1]) + " " + employee[2] + " " + employee[3] + " " + employee[4]
                
                myList = [empName, row[3], row[4], row[5], row[6], row[7], row[8]]
                #add items to row
                for item in myList:
                    qtablewidgetitem = QTableWidgetItem()
                    self.ui.attendanceTable.setItem(rowPosition, itemCount, qtablewidgetitem)
                    qtablewidgetitem.setText(str(item))
                    itemCount = itemCount+1
                rowPosition = rowPosition+1    
        else:
            dDate = self.ui.paDate.date()
            rowE = self.ui.tableWidget.currentRow()
            ID = self.ui.tableWidget.item(rowE, 0).text()
            rows = AppFunctions.getAttendance2(self, dDate, ID)
            #Create new row
            self.ui.attendanceTable2.setRowCount(0)
            for row in rows:
                #Get the number of rows
                rowPosition = self.ui.attendanceTable2.rowCount()
                if rowPosition  > row[0]:
                    continue
                itemCount = 0
                #Crete new table row
                self.ui.attendanceTable2.setRowCount(rowPosition + 1)
                qtablewidgetitem = QTableWidgetItem()
                self.ui.attendanceTable2.setVerticalHeaderItem(rowPosition, qtablewidgetitem)
                
                id = (row[1],)
                employee =  AppFunctions.getEmployeeById(id)
                
                #Name of the employee
                empName = str(employee[1]) + " " + employee[2] + " " + employee[3] + " " + employee[4]
                
                myList = [empName, row[3], row[4], row[5], row[6], row[7], row[8]]
                #add items to row
                for item in myList:
                    qtablewidgetitem = QTableWidgetItem()
                    self.ui.attendanceTable2.setItem(rowPosition, itemCount, qtablewidgetitem)
                    qtablewidgetitem.setText(str(item))
                    itemCount = itemCount+1
                rowPosition = rowPosition+1    
    
    def save(self):
        import datetime
        from docx import Document
        from docx.shared import Cm
        from docx.enum.text import WD_ALIGN_PARAGRAPH
        
        document = Document(os.path.join(os.path.dirname(__file__), 'FILE\\dtr.docx'))
        par = document.paragraphs
        
        ##Style
        style = document.styles['Normal']
        style.font.name = 'Liberation Sans Narrow'
        
        ##Name
        row = self.ui.tableWidget.currentRow()
        id = (self.ui.tableWidget.item(row, 0).text(), )
        emp = AppFunctions.getEmployeeById(id)
        par[4].text = emp[1] + " " + emp[2] + " " + emp[3] + " " + emp[4]
        par[4].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        ##Month
        date = self.ui.paDate.date()
        par[8].text = "For the month of "
        month = par[8].add_run(date.toString("MMMM"))
        par[8].add_run(", 20")
        year = par[8].add_run(date.toString("yy"))
        month.italic = month.underline = year.italic = year.underline = True
        
        table = document.tables[0]
        
        # get all attendance made by employee
        dDate = self.ui.paDate.date()
        rowE = self.ui.tableWidget.currentRow()
        ID = self.ui.tableWidget.item(rowE, 0).text()
        rs = AppFunctions.getAttendance2(self, dDate, ID)
        def convert_timedelta(duration):
            days, seconds = duration.days, duration.seconds
            hours = days * 24 + seconds // 3600
            minutes = (seconds % 3600) // 60
            seconds = (seconds % 60)
            return hours, minutes, seconds
        sum = datetime.timedelta(days=0, hours=0, minutes=0,seconds=0)
        for r in rs:
            day = int(r[2].day)
            print(day)
            row = table.rows[day+1].cells
            
            for counter in range(1,7):
                td = r[counter+2]
                dt = datetime.datetime.strptime(str(td), "%H:%M:%S")
                dts = dt.strftime('%H:%M')
                
                if counter < 5:
                    row[counter].text = str(dts)
                
                if(counter%2 == 0):
                    td2 = r[counter+1]
                    dt2 = datetime.datetime.strptime(str(td2), "%H:%M:%S")
                    time = dt-dt2
                    sum = sum + time
                    
                    if counter == 6:
                        hours, minutes, seconds = convert_timedelta(time)
                        row[counter-1].text = str(hours)
                        row[counter].text = str(minutes)
                        
        hr, min, sec = convert_timedelta(sum)
        row = table.rows[33].cells
        row[5].text = str(hr)
        row[6].text = str(min)
        
        file = filedialog.asksaveasfile(initialdir="C:",
                                        defaultextension='.docx',
                                        filetypes=[
                                            ("Microsoft Word", ".docx")
                                        ])
       
        if file:
            document.save(file.name)
        
        
    def present(id):
        from datetime import date
        from datetime import datetime as dt
        today = date.today()
        result = AppFunctions.getAttendanceWithId(id)
        t = dt.now().hour
        mysql = ""
        time = str(dt.now().hour) + ":" + str(dt.now().minute) + ":" + str(dt.now().second)  
        val = ()
        if result:
            val = (time, id, today)
            if t < 12:
                print(str(result[0][3]))
                if str(result[0][3]) != "0:00:00":
                    mysql = "UPDATE `attendance` SET `AM_timeOut` = %s WHERE employee_id = %s and date = %s;"
                else:   mysql = "UPDATE `attendance` SET `AM_timeIn` = %s WHERE employee_id = %s and date = %s;"
            elif t >= 13 and t < 15:
                print(str(result[0][5]))
                if str(result[0][5]) != "0:00:00":
                    mysql = "UPDATE `attendance` SET `PM_timeOut` = %s WHERE employee_id = %s and date = %s;"
                else:   pmysql = "UPDATE `attendance` SET `PM_timeIn` = %s WHERE employee_id = %s and date = %s;"
            elif t > 15:
                print(str(result[0][7]))
                if str(result[0][7]) != "0:00:00":
                    mysql = "UPDATE `attendance` SET `OT_timeOut` = %s WHERE employee_id = %s and date = %s;"
                else:   mysql = "UPDATE `attendance` SET `OT_timeIn` = %s WHERE employee_id = %s and date = %s;"
        else:
            val = (id, today, time, time)
            if t < 12:
                mysql = """INSERT INTO attendance (`id`, `employee_id`, `date`, `AM_timeIn`, `AM_timeOut`, `PM_timeIn`,
                `PM_timeOut`, `OT_timeIn`, `OT_timeOut`) VALUES (NULL, %s, %s, %s, %s, '', '', '', '');"""
            elif t >= 13 and t < 15:
                mysql = """INSERT INTO attendance (`id`, `employee_id`, `date`, `AM_timeIn`, `AM_timeOut`, `PM_timeIn`,
                `PM_timeOut`, `OT_timeIn`, `OT_timeOut`) VALUES (NULL, %s, %s, '', '', %s, %s, '', '');"""
            elif t > 15:
                mysql = """INSERT INTO attendance (`id`, `employee_id`, `date`, `AM_timeIn`, `AM_timeOut`, `PM_timeIn`,
                `PM_timeOut`, `OT_timeIn`, `OT_timeOut`) VALUES (NULL, %s, %s, '', '', '', '', %s, %s);"""

        conn = AppFunctions.create_connection()
        mycursor = conn.cursor()
        mycursor.execute(mysql, val)
        conn.commit()
        